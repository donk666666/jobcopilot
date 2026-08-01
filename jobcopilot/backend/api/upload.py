"""
简历文件上传与解析 API
"""

from fastapi import APIRouter, HTTPException, UploadFile, File

router = APIRouter(prefix="/api/resume", tags=["简历上传"])

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
ALLOWED_EXTENSIONS = {".docx", ".pdf"}


@router.post("/upload")
async def upload_resume(file: UploadFile = File(...)):
    """上传简历文件（.docx / .pdf），自动解析提取文本"""

    # 校验文件扩展名
    filename = file.filename or ""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    ext_with_dot = f".{ext}"
    if ext_with_dot not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的格式 '.{ext}'，请上传 .docx 或 .pdf 文件",
        )

    # 读取文件内容
    content = await file.read()

    # 校验大小
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"文件过大（{len(content) / 1024 / 1024:.1f}MB），最大支持 5MB",
        )

    if len(content) == 0:
        raise HTTPException(status_code=400, detail="文件内容为空，请重新选择")

    # 根据类型解析
    try:
        if ext == "docx":
            text = _parse_docx(content)
        else:
            text = _parse_pdf(content)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=400, detail=f"文件解析失败，请确认文件未损坏: {str(e)}"
        )

    if not text.strip():
        raise HTTPException(status_code=400, detail="未能从文件中提取到文字内容，请确认文件包含可读文本")

    return {
        "filename": filename,
        "text": text.strip(),
        "char_count": len(text.strip()),
        "file_type": "docx" if ext == "docx" else "pdf",
    }


def _parse_docx(content: bytes) -> str:
    """解析 .docx 文件，提取段落和表格文本"""
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(content))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


def _parse_pdf(content: bytes) -> str:
    """解析 PDF 文件，逐页提取文本"""
    from io import BytesIO
    import pdfplumber

    parts = []
    with pdfplumber.open(BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text.strip())

    return "\n".join(parts)
