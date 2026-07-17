import docx
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def font(run, name='微软雅黑', size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def add_title(text, size=24):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    font(r, size=size, bold=True, color=RGBColor(0x1a, 0x1a, 0x1a))
    return p

def add_subtitle(text, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    font(r, size=size, color=RGBColor(0x88, 0x88, 0x88))
    return p

def add_section_header(text):
    """带左侧蓝色色块 + 右侧分割线的美观标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0

    # 左侧蓝色竖条用 ▌字符模拟
    r_bar = p.add_run('▌ ')
    font(r_bar, size=12, bold=True, color=RGBColor(0x1a, 0x56, 0xdb))

    r_text = p.add_run(text)
    font(r_text, size=13, bold=True, color=RGBColor(0x1a, 0x1a, 0x1a))

    # 右侧灰色分割线
    r_line = p.add_run(' ' + '─' * 60)
    font(r_line, size=6, color=RGBColor(0xcc, 0xcc, 0xcc))

    return p

def add_info_row(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    r1 = p.add_run(f'{label}：')
    font(r1, size=10.5, bold=True, color=RGBColor(0x33, 0x33, 0x33))
    r2 = p.add_run(value)
    font(r2, size=10.5)

def add_bullet(text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run('· ' + text)
    font(r, size=size)
    return p

def add_project_header(text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run(text)
    font(r, size=size, bold=True, color=RGBColor(0x1a, 0x1a, 0x1a))
    return p

def add_normal(text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run(text)
    font(r, size=size)
    return p

def add_sub_header(text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(text)
    font(r, size=size, bold=True, color=RGBColor(0x44, 0x44, 0x44))
    return p

# ===== HEADER =====
add_title('萧 仁 科')
add_subtitle('求职意向：AI 项目实习生')

# ===== 个人信息 =====
add_section_header('个人信息')
add_info_row('电话', '(+86) 13728774722')
add_info_row('邮箱', '1426454082@qq.com')
add_info_row('微信', 'Xiaopro556677')
add_info_row('性别', '男')
add_info_row('年龄', '21 岁')
add_info_row('民族', '汉')
add_info_row('籍贯', '深圳')

# ===== 教育经历 =====
add_section_header('教育经历')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
p.paragraph_format.line_spacing = 1.35
r1 = p.add_run('华南理工大学')
font(r1, size=10.5, bold=True)
r2 = p.add_run('  ·  大数据管理与应用  ·  本科  |  2023.9 - 2027.6')
font(r2, size=10.5)
add_bullet('GPA 3.65 / 4.0（前 20%），英语四六级，校三等奖学金')

# ===== 实习与项目经历 =====
add_section_header('实习与项目经历')

# --- 1. AutoGEO ---
add_project_header('广州架构矩阵有限公司  ·  AI 产品实习生  |  2026.1 - 2026.3')
add_normal('主导 AutoGEO 产品中 n8n AI 工作流引擎的设计与开发。AutoGEO 基于 Electron + Vue 3 + FastAPI + n8n 四层架构，集成 DeepSeek 实现账号管理、内容生成、发布管理及浏览器自动化。')
add_sub_header('n8n 工作流贡献')
add_bullet('关键词蒸馏：大模型自动提炼长尾关键词，规则提取与合并生成优化策略，辅助 GEO 内容定向')
add_bullet('AI 收录检测：千问 / DeepSeek 多引擎自动收录检测管线，定时调度与趋势分析预警')
add_bullet('GEO 文章生成：DeepSeek API 异步文章生成，支持 Webhook 回调、字数控制与多平台适配')

# --- 2. JobCopilot ---
add_project_header('JobCopilot AI 求职助手（个人全栈项目）  |  2026.6')
add_normal('基于 LangChain + RAG + Prompt Engineering 的智能求职平台，Vue 3 + FastAPI + SQLite + ChromaDB 全栈，实现 JD 解析、简历匹配、求职信生成与投递管理全流程 AI 辅助。')
add_sub_header('核心实现')
add_bullet('AI Agent 引擎：LangChain ReAct 模式 Thought → Action → Observation 循环，集成 4 个自定义工具（JD 分析器 / 简历匹配器 / 简历改写器 / 求职信生成器），支持多步骤任务编排与流式输出')
add_bullet('RAG 检索增强：ChromaDB + BGE-small-zh-v1.5 双集合向量库（简历库 / JD 库），RecursiveCharacterTextSplitter 智能分块，跨库语义检索')
add_bullet('Prompt Engineering：Few-shot / Chain-of-Thought / Prompt Chaining 三种策略，JSON Schema 约束输出一致性')
add_bullet('AI 效果评估：bad case 分析、输出质量验证，沉淀可复用的 prompt 测试与优化迭代模板')

# --- 3. 智能文档问答助手 ---
add_project_header('智能文档问答助手（个人全栈项目）  |  2026.7')
add_normal('基于 LangGraph + ChromaDB + GLM-5.2 + FastAPI 构建的 RAG 智能问答系统，支持多轮对话、文档入库、RSS 自动抓取及飞书机器人接入。已部署至腾讯云轻量服务器对外提供服务。')
add_sub_header('核心实现')
add_bullet('LangGraph 工作流：设计 rewrite_query → retrieve → judge → generate 四节点 Agent 流程，砍掉冗余的意图分类与澄清节点，单次问答 LLM 调用从 4 次降至 2 次，延迟减半')
add_bullet('RAG 知识库：ChromaDB + BGE-small-zh-v1.5 构建向量库（512d），RecursiveCharacterTextSplitter（500/80）中文友好分块；向量相似度 0.6 + 关键词匹配 0.4 混合检索')
add_bullet('飞书机器人：lark-oapi 实现 /feishu/callback 消息回调，支持私聊与群聊 @ 问答，Session 级对话记忆')
add_bullet('生产部署：Docker Compose 一键部署，含 HEALTHCHECK 健康检查、路径穿越防护与文件类型白名单、腾讯云镜像加速')
add_sub_header('项目亮点')
add_bullet('Embedding 模型本地离线部署于 Docker 镜像内，无需访问 HuggingFace，开箱即用')
add_bullet('集成 RSS 定时抓取（阮一峰 / 美团技术 / InfoQ），自动入库 ChromaDB，知识库可持续增长')
add_bullet('Docker 单容器覆盖 FastAPI + ChromaDB + Embedding + RSS 调度 + 静态前端全部服务')

# --- 竞赛 ---
add_project_header('竞赛经历')
add_bullet('美国数学建模大赛（MCM/ICM）、黑客马拉松、大学生创新创业训练计划（大创）、三创赛等')

# ===== 个人技能 =====
add_section_header('个人技能')
add_bullet('熟练掌握 Python、FastAPI、LangChain / LangGraph、ChromaDB 向量数据库与 RAG 检索增强技术')
add_bullet('熟练运用 Prompt Engineering 与 AI Agent 开发；具备 AI 应用效果评估、bad case 分析与 prompt 优化经验')
add_bullet('掌握 Docker 容器化部署、MySQL、Linux 常用运维，有云服务器从零部署上线的完整实践经验')

output = r'C:\Users\14264\Desktop\新项目\萧仁科简历.docx'
doc.save(output)
print('Done!')
